import os
import shutil
import glob
import re
from datetime import datetime
from typing import List, Dict, Any, Optional
import docx
from docx.shared import Inches
import pandas as pd
import json
import yaml
import markdown
import pdfkit
from PIL import Image
import mimetypes

class FileManager:
    def __init__(self):
        self.supported_doc_types = {
            'text': ['.txt', '.md'],
            'document': ['.docx', '.doc'],
            'spreadsheet': ['.xlsx', '.xls', '.csv'],
            'presentation': ['.pptx', '.ppt'],
            'pdf': ['.pdf'],
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp'],
            'audio': ['.mp3', '.wav', '.ogg'],
            'video': ['.mp4', '.avi', '.mov', '.wmv']
        }

    def create_file(self, file_path: str, content: str = "", file_type: str = None) -> str:
        """
        Create a new file with optional content.
        Args:
            file_path: Path where the file should be created
            content: Initial content for the file
            file_type: Type of file to create (if not specified, inferred from extension)
        """
        try:
            # Create directory if it doesn't exist
            os.makedirs(os.path.dirname(os.path.abspath(file_path)), exist_ok=True)
            
            # Determine file type if not specified
            if not file_type:
                file_type = os.path.splitext(file_path)[1].lower()
            
            # Create file based on type
            if file_type in self.supported_doc_types['document']:
                doc = docx.Document()
                doc.add_paragraph(content)
                doc.save(file_path)
            elif file_type in self.supported_doc_types['spreadsheet']:
                df = pd.DataFrame()
                if content:
                    # Try to parse content as JSON or CSV
                    try:
                        data = json.loads(content)
                        df = pd.DataFrame(data)
                    except:
                        # If not JSON, treat as CSV
                        df = pd.read_csv(pd.StringIO(content))
                df.to_excel(file_path, index=False)
            elif file_type in self.supported_doc_types['pdf']:
                # Convert markdown to PDF
                html = markdown.markdown(content)
                pdfkit.from_string(html, file_path)
            else:
                # Default to text file
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            return f"File created successfully at {file_path}"
        except Exception as e:
            return f"Error creating file: {str(e)}"

    def delete_file(self, file_path: str) -> str:
        """
        Delete a file or directory.
        Args:
            file_path: Path to the file or directory to delete
        """
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
                return f"File {file_path} deleted successfully"
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
                return f"Directory {file_path} deleted successfully"
            else:
                return f"Path {file_path} does not exist"
        except Exception as e:
            return f"Error deleting file/directory: {str(e)}"

    def search_files(self, pattern: str, directory: str = ".", recursive: bool = True) -> List[str]:
        """
        Search for files matching a pattern.
        Args:
            pattern: File pattern to search for (supports glob patterns)
            directory: Directory to search in
            recursive: Whether to search recursively
        """
        try:
            if recursive:
                return glob.glob(os.path.join(directory, "**", pattern), recursive=True)
            return glob.glob(os.path.join(directory, pattern))
        except Exception as e:
            return [f"Error searching files: {str(e)}"]

    def organize_files(self, directory: str, organization_type: str = "extension") -> str:
        """
        Organize files in a directory based on specified criteria.
        Args:
            directory: Directory to organize
            organization_type: Type of organization ('extension', 'date', 'type')
        """
        try:
            if not os.path.exists(directory):
                return f"Directory {directory} does not exist"

            for filename in os.listdir(directory):
                file_path = os.path.join(directory, filename)
                if os.path.isfile(file_path):
                    if organization_type == "extension":
                        ext = os.path.splitext(filename)[1][1:].lower()
                        if not ext:
                            ext = "no_extension"
                        target_dir = os.path.join(directory, ext)
                    elif organization_type == "date":
                        mod_time = datetime.fromtimestamp(os.path.getmtime(file_path))
                        target_dir = os.path.join(directory, mod_time.strftime("%Y-%m-%d"))
                    elif organization_type == "type":
                        mime_type = mimetypes.guess_type(file_path)[0]
                        if mime_type:
                            type_dir = mime_type.split('/')[0]
                        else:
                            type_dir = "unknown"
                        target_dir = os.path.join(directory, type_dir)
                    else:
                        return f"Invalid organization type: {organization_type}"

                    os.makedirs(target_dir, exist_ok=True)
                    shutil.move(file_path, os.path.join(target_dir, filename))

            return f"Files organized successfully in {directory}"
        except Exception as e:
            return f"Error organizing files: {str(e)}"

    def batch_rename(self, directory: str, pattern: str, new_pattern: str) -> str:
        """
        Rename multiple files matching a pattern.
        Args:
            directory: Directory containing files to rename
            pattern: Pattern to match files (supports regex)
            new_pattern: New pattern for file names (supports {n} for numbering)
        """
        try:
            files = self.search_files(pattern, directory)
            for i, file_path in enumerate(files, 1):
                dir_name = os.path.dirname(file_path)
                old_name = os.path.basename(file_path)
                new_name = new_pattern.format(n=i)
                new_path = os.path.join(dir_name, new_name)
                os.rename(file_path, new_path)
            return f"Renamed {len(files)} files successfully"
        except Exception as e:
            return f"Error renaming files: {str(e)}"

    def convert_file(self, input_path: str, output_path: str) -> str:
        """
        Convert file from one format to another.
        Args:
            input_path: Path to input file
            output_path: Path for output file
        """
        try:
            input_ext = os.path.splitext(input_path)[1].lower()
            output_ext = os.path.splitext(output_path)[1].lower()

            if input_ext in self.supported_doc_types['document'] and output_ext == '.pdf':
                doc = docx.Document(input_path)
                html = ""
                for para in doc.paragraphs:
                    html += f"<p>{para.text}</p>"
                pdfkit.from_string(html, output_path)
            elif input_ext in ['.jpg', '.jpeg', '.png'] and output_ext in ['.jpg', '.jpeg', '.png']:
                img = Image.open(input_path)
                img.save(output_path)
            elif input_ext == '.md' and output_ext == '.html':
                with open(input_path, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                html_content = markdown.markdown(md_content)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
            else:
                return f"Unsupported conversion from {input_ext} to {output_ext}"

            return f"File converted successfully to {output_path}"
        except Exception as e:
            return f"Error converting file: {str(e)}"

    def create_document(self, doc_type: str, content: Dict[str, Any], output_path: str) -> str:
        """
        Create a document of specified type.
        Args:
            doc_type: Type of document ('text', 'document', 'spreadsheet', 'presentation', 'pdf')
            content: Content for the document
            output_path: Path where to save the document
        """
        try:
            if doc_type == 'document':
                doc = docx.Document()
                for section in content.get('sections', []):
                    doc.add_heading(section.get('title', ''), level=1)
                    for para in section.get('paragraphs', []):
                        doc.add_paragraph(para)
                doc.save(output_path)
            elif doc_type == 'spreadsheet':
                df = pd.DataFrame(content.get('data', []))
                df.to_excel(output_path, index=False)
            elif doc_type == 'presentation':
                # Implementation for presentation creation
                pass
            elif doc_type == 'pdf':
                html_content = ""
                for section in content.get('sections', []):
                    html_content += f"<h1>{section.get('title', '')}</h1>"
                    for para in section.get('paragraphs', []):
                        html_content += f"<p>{para}</p>"
                pdfkit.from_string(html_content, output_path)
            else:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content.get('text', ''))

            return f"{doc_type.capitalize()} document created successfully at {output_path}"
        except Exception as e:
            return f"Error creating document: {str(e)}"

# Create a singleton instance
file_manager = FileManager() 